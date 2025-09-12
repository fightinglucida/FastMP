import os
import re
import time
import threading
import requests
import traceback
import logging
import zipfile
import pypandoc
from queue import Queue, Empty
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from datetime import datetime

class WeChatArticleDownloader:
    def _clean_filename(self, filename):
        """清理文件名中的非法字符"""
        return ''.join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).rstrip()
    """微信文章下载器，负责下载单篇文章"""
    
    def __init__(self, save_dir="."):
        self.save_dir = save_dir
        self.images_dir = os.path.join(save_dir, "images")
        os.makedirs(self.images_dir, exist_ok=True)
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
        }
        # 当前文章标题
        self.current_article_title = None
        # 当前文章发布时间
        self.current_publish_time = ""
        # 设置日志
        self._setup_logger()
        
    def _setup_logger(self):
        """设置日志记录器"""
        self.logger = logging.getLogger("WeChatArticleDownloader")
        self.logger.setLevel(logging.INFO)
        
        # 清除所有已存在的处理器
        for handler in self.logger.handlers[:]:
            self.logger.removeHandler(handler)
        
        # 创建控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        
        # 设置日志格式
        formatter = logging.Formatter('%(message)s')
        console_handler.setFormatter(formatter)
        
        # 添加处理器到日志记录器
        self.logger.addHandler(console_handler)
        
    def download_article(self, url, file_format='md', account_name=None):
        """下载文章并保存为指定格式
        
        Args:
            url: 文章链接
            file_format: 'md'/'docx'/'pdf'
            account_name: 公众号名（用于保存路径）
            
        Returns:
            tuple: (是否成功, 文件路径)
        """
        try:
            # 获取文章内容
            title, content = self.get_article_content(url)
            
            if title and content:
                # 生成保存路径
                if account_name:
                    base_dir = os.path.join('static', 'mp_accounts', account_name)
                else:
                    base_dir = self.save_dir
                if file_format == 'md':
                    save_dir = base_dir
                elif file_format == 'docx':
                    save_dir = os.path.join(base_dir, 'docx')
                elif file_format == 'pdf':
                    save_dir = os.path.join(base_dir, 'pdf')
                else:
                    save_dir = base_dir
                os.makedirs(save_dir, exist_ok=True)
                self.save_dir = save_dir
                if file_format == 'md':
                    file_path = self.save_to_markdown(title, content)
                elif file_format == 'docx':
                    # 用 python-docx 方案
                    html_content = requests.get(url, headers=self.headers, timeout=30).text
                    file_path = self.save_to_docx_with_python_docx(title, html_content)
                elif file_format == 'pdf':
                    file_path = self.save_to_pdf(title, content)
                else:
                    file_path = self.save_to_markdown(title, content)
                if file_path:
                    return True, file_path
            
            return False, None
            
        except Exception as e:
            self.logger.error(f"下载文章失败: {str(e)}")
            return False, None

    def get_article_content(self, url):
        """获取文章内容并下载图片
        
        Args:
            url (str): 微信公众号文章URL
            
        Returns:
            tuple: (标题, Markdown内容) 或 (None, None)
        """
        try:
            # 获取页面内容
            response = requests.get(url, headers=self.headers, timeout=30)
            response.encoding = 'utf-8'
            
            if response.status_code != 200:
                self.logger.error(f"请求失败，状态码: {response.status_code}")
                return None, None
                
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 获取文章标题
            title_element = soup.find('h1', class_='rich_media_title')
            if not title_element:
                title_element = soup.find('h1', id='activity-name')
            
            if not title_element:
                # 尝试从其他地方获取标题
                title_match = re.search(r'var msg_title = "([^"]+)"', response.text)
                if title_match:
                    title = title_match.group(1)
                else:
                    self.logger.error("无法找到文章标题")
                    return None, None
            else:
                title = title_element.get_text().strip()
            
            # 获取发布时间
            publish_time_str = self._extract_publish_time(response.text, soup)
            
            # 设置当前文章标题和时间戳
            self.current_article_title = title
            self.current_publish_time = publish_time_str
            self.logger.info(f"[下载中]：{title}，时间戳：{publish_time_str}")
            
            # 获取文章内容
            content_element = soup.find('div', id='js_content')
            if not content_element:
                content_element = soup.find('div', class_='rich_media_content')
            
            if not content_element:
                content_element = soup.find('div', class_='js_underline_content')
                
            if not content_element:
                content_element = soup.find('div', class_=lambda c: c and ('rich_media_content' in c or 'js_underline_content' in c))
                
            if not content_element:
                self.logger.error("无法找到文章内容区域")
                return None, None
            
            # 预处理内容元素，保持原始结构
            self._preprocess_content(content_element)
            
            # 处理图片
            img_index = 0
            for img in content_element.find_all('img'):
                # 获取图片URL
                img_url = img.get('data-src') or img.get('src')
                if img_url:
                    img_url = urljoin(url, img_url)
                    
                    # 下载图片，使用序号
                    local_filename = self.download_image(img_url, img_index)
                    if local_filename:
                        # 更新图片链接为本地路径
                        img['src'] = f'./images/{local_filename}'
                        if 'data-src' in img.attrs:
                            del img['data-src']
                    img_index += 1  # 始终递增 index，无论是否下载成功
            
            # 转换为Markdown格式
            markdown_content = self._convert_to_markdown(title, content_element)
            
            return title, markdown_content
            
        except Exception as e:
            self.logger.error(f"获取文章内容失败: {str(e)}")
            traceback.print_exc()
            return None, None
    
    def _extract_publish_time(self, html_text, soup):
        """提取发布时间"""
        publish_time_str = ""
        
        # 方法1: 直接从HTML源码中搜索日期格式
        date_patterns = [
            r'(\d{4})年(\d{1,2})月(\d{1,2})日\s+(\d{1,2}):(\d{1,2})',
            r'(\d{4})-(\d{1,2})-(\d{1,2})\s+(\d{1,2}):(\d{1,2})',
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, html_text)
            if matches:
                year, month, day, hour, minute = matches[0]
                month = month.zfill(2) if len(month) == 1 else month
                day = day.zfill(2) if len(day) == 1 else day
                hour = hour.zfill(2) if len(hour) == 1 else hour
                minute = minute.zfill(2) if len(minute) == 1 else minute
                publish_time_str = f"{year}{month}{day}{hour}{minute}"
                break
        
        # 方法2: 查找publish_time元素
        if not publish_time_str:
            publish_time_element = soup.find('em', id='publish_time')
            if publish_time_element:
                publish_time_text = publish_time_element.get_text().strip()
                if publish_time_text:
                    time_match = re.search(r'(\d{4})年(\d{2})月(\d{2})日\s+(\d{2}):(\d{2})', publish_time_text)
                    if time_match:
                        year, month, day, hour, minute = time_match.groups()
                        publish_time_str = f"{year}{month}{day}{hour}{minute}"
        
        # 如果都失败，使用当前时间
        if not publish_time_str:
            now = datetime.now()
            publish_time_str = now.strftime("%Y%m%d%H%M")
            
        return publish_time_str
    
    def _preprocess_content(self, content_element):
        """预处理内容元素"""
        # 处理所有的data-src属性
        for img in content_element.find_all('img'):
            if img.get('data-src') and not img.get('src'):
                img['src'] = img['data-src']
        
        # 处理所有的<br>标签
        for br in content_element.find_all('br'):
            br.replace_with('\n')
        
        # 处理所有的<section>标签
        for section in content_element.find_all('section'):
            if not section.find(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'blockquote']):
                section.name = 'p'
    
    def _convert_to_markdown(self, title, content_element):
        """将HTML内容转换为Markdown格式"""
        markdown_content = f'# {title}\n\n'
        
        # 获取所有顶级元素
        top_elements = list(content_element.children)
        top_elements = [el for el in top_elements if not (isinstance(el, str) and el.strip() == '')]
        
        # 处理每个顶级元素
        for element in top_elements:
            if isinstance(element, str):
                text = element.strip()
                if text:
                    markdown_content += text + '\n\n'
                continue
            
            # 处理图片节点
            if element.name == 'img':
                img_src = element.get('src')
                if img_src and (img_src.startswith('./images/') or img_src.startswith('images/')):
                    # 修正图片路径为 images/xxx.jpg
                    img_src = img_src.replace('./images/', 'images/')
                    markdown_content += f'![图片]({img_src})\n\n'
                continue
            
            # 处理链接节点
            elif element.name == 'a':
                href = element.get('href', '')
                link_text = element.get_text().strip()
                if link_text and href:
                    markdown_content += f'[{link_text}]({href})\n\n'
                elif link_text:
                    markdown_content += f'{link_text}\n\n'
                continue
            
            # 处理段落
            if element.name == 'p':
                images = element.find_all('img')
                if images and len(list(element.stripped_strings)) == 0:
                    for img in images:
                        img_src = img.get('src')
                        if img_src and (img_src.startswith('./images/') or img_src.startswith('images/')):
                            img_src = img_src.replace('./images/', 'images/')
                            markdown_content += f'![图片]({img_src})\n\n'
                else:
                    p_content = self._process_inline_elements(element)
                    if p_content.strip():
                        markdown_content += p_content + '\n\n'
            
            # 处理标题
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                text = element.get_text().strip()
                if text:
                    markdown_content += f'{"#" * level} {text}\n\n'
            
            # 处理其他元素
            elif element.name in ['div', 'section']:
                div_content = self._process_div_or_section(element)
                if div_content:
                    markdown_content += div_content
        
        return markdown_content.strip()
    
    def _process_inline_elements(self, element):
        """处理段落中的内联元素，递归优先识别链接"""
        result = ""
        for child in element.children:
            if isinstance(child, str):
                result += child
            elif child.name == 'br':
                result += '\n'
            elif child.name == 'img':
                img_src = child.get('src')
                if img_src and (img_src.startswith('./images/') or img_src.startswith('images/')):
                    img_src = img_src.replace('./images/', 'images/')
                    result += f'![图片]({img_src})'
            elif child.name in ['strong', 'b']:
                result += f"**{self._process_inline_elements(child)}**"
            elif child.name in ['em', 'i']:
                result += f"*{self._process_inline_elements(child)}*"
            elif child.name == 'a':
                href = child.get('href', '')
                link_text = self._process_inline_elements(child)
                if link_text and href:
                    result += f"[{link_text}]({href})"
                elif link_text:
                    result += link_text
            elif child.name == 'code':
                result += f"`{child.get_text().strip()}`"
            else:
                # 对于 span、div、section 等，递归处理所有子节点
                result += self._process_inline_elements(child)
        result = re.sub(r' +', ' ', result)
        result = result.replace('\n', '\n\n')
        return result.strip()

    def _process_div_or_section(self, element):
        """递归处理div或section元素，优先识别链接"""
        result = ""
        images = element.find_all('img', recursive=False)
        if len(images) == 1 and len(list(element.stripped_strings)) == 0:
            img = images[0]
            img_src = img.get('src')
            if img_src and (img_src.startswith('./images/') or img_src.startswith('images/')):
                img_src = img_src.replace('./images/', 'images/')
                return f'![图片]({img_src})\n\n'
        for child in element.children:
            if isinstance(child, str):
                text = child.strip()
                if text:
                    result += text + '\n\n'
            elif child.name == 'img':
                img_src = child.get('src')
                if img_src and (img_src.startswith('./images/') or img_src.startswith('images/')):
                    img_src = img_src.replace('./images/', 'images/')
                    result += f'![图片]({img_src})\n\n'
            elif child.name == 'p':
                p_content = self._process_inline_elements(child)
                if p_content:
                    result += p_content + '\n\n'
            elif child.name == 'a':
                href = child.get('href', '')
                link_text = self._process_inline_elements(child)
                if link_text and href:
                    result += f'[{link_text}]({href})\n\n'
                elif link_text:
                    result += f'{link_text}\n\n'
            elif child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(child.name[1])
                text = child.get_text().strip()
                if text:
                    result += f'{"#" * level} {text}\n\n'
            elif child.name in ['div', 'section', 'span']:
                div_content = self._process_div_or_section(child)
                if div_content:
                    result += div_content
            else:
                # 其他元素递归处理
                result += self._process_inline_elements(child)
        return result
    
    def download_image(self, img_url, index=None):
        """下载图片到指定目录"""
        try:
            response = requests.get(img_url, stream=True, headers=self.headers, timeout=30)
            
            if response.status_code == 200:
                # 使用文章标题和序号生成文件名
                if self.current_article_title and index is not None:
                    safe_title = re.sub(r'[<>:"/\\|?*]', '_', self.current_article_title)
                    if self.current_publish_time:
                        filename = f'{self.current_publish_time}-{safe_title}_{index:03d}.jpg'
                    else:
                        filename = f'{safe_title}_{index:03d}.jpg'
                else:
                    filename = os.path.basename(img_url.split('?')[0])
                    if not filename or len(filename) > 100 or not filename.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
                        filename = f"img_{int(time.time() * 1000)}_{hash(img_url) % 10000}.jpg"
                
                # 确保文件名是合法的
                filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                
                # 保存图片
                file_path = os.path.join(self.images_dir, filename)
                with open(file_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                
                return filename
            else:
                self.logger.error(f"下载图片失败，状态码: {response.status_code}, URL: {img_url}")
                return None
        except Exception as e:
            self.logger.error(f"下载图片失败 {img_url}: {str(e)}")
            return None
    
    def save_to_markdown(self, title, content):
        """保存内容为Markdown文件"""
        try:
            # 清理文件名中的非法字符
            safe_title = self._clean_filename(title)
            
            # 如果有发布时间，添加到文件名前
            if self.current_publish_time:
                filename = f"{self.current_publish_time}-{safe_title}.md"
            else:
                filename = f"{safe_title}.md"
                
            file_path = os.path.join(self.save_dir, filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            self.logger.info(f"[已保存]：{file_path}")
            return file_path
        except Exception as e:
            self.logger.error(f"保存Markdown文件失败: {str(e)}")
            return None
    
    def save_to_docx(self, title, markdown_content):
        """保存内容为docx文件"""
        try:
            safe_title = self._clean_filename(title)
            if self.current_publish_time:
                filename = f"{self.current_publish_time}-{safe_title}.docx"
            else:
                filename = f"{safe_title}.docx"
            file_path = os.path.join(self.save_dir, filename)
            # 指定图片资源路径
            images_path = os.path.abspath(os.path.join(self.save_dir, "..", "images"))
            pypandoc.convert_text(
                markdown_content,
                'docx',
                format='md',
                outputfile=file_path,
                extra_args=[f'--resource-path={images_path}']
            )
            self.logger.info(f"[已保存]：{file_path}")
            return file_path
        except Exception as e:
            self.logger.error(f"保存docx文件失败: {str(e)}")
            return None

    def save_to_pdf(self, title, markdown_content):
        """保存内容为pdf文件"""
        try:
            safe_title = self._clean_filename(title)
            if self.current_publish_time:
                filename = f"{self.current_publish_time}-{safe_title}.pdf"
            else:
                filename = f"{safe_title}.pdf"
            file_path = os.path.join(self.save_dir, filename)
            pypandoc.convert_text(markdown_content, 'pdf', format='md', outputfile=file_path)
            self.logger.info(f"[已保存]：{file_path}")
            return file_path
        except Exception as e:
            self.logger.error(f"保存pdf文件失败: {str(e)}")
            return None
    
    def save_to_docx_with_python_docx(self, title, html_content):
        """修复图片重复插入问题，保证每张图片只插入一次"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        safe_title = self._clean_filename(title)
        if self.current_publish_time:
            filename = f"{self.current_publish_time}-{safe_title}.docx"
        else:
            filename = f"{safe_title}.docx"
        file_path = os.path.join(self.save_dir, filename)
        doc = Document()
        doc.add_heading(title, level=1)
        soup = BeautifulSoup(html_content, 'html.parser')
        content = soup.find('div', id='js_content')
        if not content:
            content = soup.find('div', class_='rich_media_content')
        if not content:
            content = soup.find('div', class_='js_underline_content')
        if not content:
            content = soup
        
        # 收集已下载图片（使用实际文件名模式匹配）
        img_files = []
        # 尝试多种可能的文件名模式
        patterns = [
            f"{self.current_publish_time}-{self.current_article_title}_",
            f"{self.current_publish_time}-{safe_title}_",
            f"{self.current_publish_time}-",
        ]
        
        self.logger.info(f"images_dir 中的所有文件: {os.listdir(self.images_dir)}")
        
        for pattern in patterns:
            pattern_clean = re.sub(r'[<>:"/\\|?*]', '_', pattern)
            self.logger.info(f"尝试匹配模式: {pattern_clean}")
            
            for fname in sorted(os.listdir(self.images_dir)):
                if fname.startswith(pattern_clean) and fname.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
                    img_path = os.path.join(self.images_dir, fname)
                    if img_path not in img_files:  # 避免重复
                        img_files.append(img_path)
                        self.logger.info(f"找到图片: {fname}")
            
            if img_files:  # 如果找到了图片就停止尝试其他模式
                break
        
        self.logger.info(f"共找到 {len(img_files)} 张图片")
        img_index = 0
        
        def get_full_text(elem):
            # 递归获取所有子节点文本
            if isinstance(elem, str):
                return elem.strip()
            text = ''
            for child in elem.children:
                text += get_full_text(child)
            return text
        
        def add_hyperlink(paragraph, url, text):
            # 通过字段代码插入超链接
            try:
                part = paragraph.part
                r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                # 超链接样式
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rPr.append(rStyle)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = text
                new_run.append(t)
                hyperlink.append(new_run)
                paragraph._p.append(hyperlink)
                self.logger.info(f"成功插入超链接: {text} -> {url}")
            except Exception as e:
                self.logger.error(f"插入超链接失败: {text} -> {url}, 错误: {e}")
                # 如果超链接插入失败，至少插入文本
                paragraph.add_run(text)
        
        def add_element(elem, parent_paragraph=None):
            nonlocal img_index
            if elem is None:
                return
            
            # 优先处理链接标签
            if elem.name == 'a':
                url = elem.get('href', '')
                link_text = get_full_text(elem)
                self.logger.info(f"发现链接: 文本='{link_text}', URL='{url}'")
                if parent_paragraph and url and link_text:
                    add_hyperlink(parent_paragraph, url, link_text)
                elif parent_paragraph and link_text:
                    parent_paragraph.add_run(link_text)
                return
            if isinstance(elem, str):
                text = elem.strip()
                if text:
                    if parent_paragraph:
                        parent_paragraph.add_run(text)
                    else:
                        p = doc.add_paragraph(text)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif elem.name == 'img':
                if img_index < len(img_files):
                    try:
                        self.logger.info(f"插入图片 {img_index + 1}: {img_files[img_index]}")
                        doc.add_picture(img_files[img_index], width=Inches(4))
                        img_index += 1
                    except Exception as e:
                        self.logger.error(f"插入图片失败: {img_files[img_index]}, {e}")
                        img_index += 1
                else:
                    self.logger.warning(f"图片索引 {img_index} 超出范围，共有 {len(img_files)} 张图片")
            elif elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                text = elem.get_text(strip=True)
                if text:
                    level = int(elem.name[1]) if elem.name[1].isdigit() else 2
                    doc.add_heading(text, level=level)
            elif elem.name == 'p':
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for child in elem.children:
                    add_element(child, parent_paragraph=p)
            elif elem.name == 'br':
                if parent_paragraph:
                    parent_paragraph.add_run('\n')
            elif elem.name in ['strong', 'b', 'em', 'i', 'span', 'div', 'section']:
                # 递归处理所有子元素，优先查找链接
                for child in elem.children:
                    add_element(child, parent_paragraph=parent_paragraph)
            elif elem.name == 'hr':
                doc.add_paragraph('------------------------')
            elif elem.name in ['ul', 'ol']:
                for li in elem.find_all('li', recursive=False):
                    # 对列表项也要递归处理，可能包含链接
                    li_p = doc.add_paragraph(style='List Bullet' if elem.name=='ul' else 'List Number')
                    li_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for li_child in li.children:
                        add_element(li_child, parent_paragraph=li_p)
            elif elem.name == 'blockquote':
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.style.font.italic = True
                # 递归处理引用内容，可能包含链接
                for child in elem.children:
                    add_element(child, parent_paragraph=p)
            else:
                # 对于其他未知元素，也要递归处理
                for child in elem.children:
                    add_element(child, parent_paragraph=parent_paragraph)
        for elem in content.children:
            add_element(elem)
        
        doc.save(file_path)
        self.logger.info(f"[已保存]：{file_path}，插入了 {img_index} 张图片")
        return file_path


class ArticleDownloadManager:
    """文章下载管理器，管理多线程下载"""
    
    def __init__(self, save_dir=".", max_threads=5):
        self.save_dir = save_dir
        self.download_queue = Queue()
        self.download_threads = []
        self.is_downloading = False
        self.download_results = {}
        self.max_threads = max_threads
        self.progress_callback = None
        
    def add_article(self, article_info):
        """添加文章到下载队列"""
        self.download_queue.put(article_info)
        self.download_results[article_info['url']] = {'status': '等待下载', 'file_path': None}
        
    def start_download(self, progress_callback=None):
        """开始下载队列中的文章"""
        if self.is_downloading:
            return
            
        self.progress_callback = progress_callback
        self.is_downloading = True
        self.download_threads = []
        
        for _ in range(self.max_threads):
            thread = threading.Thread(target=self._download_worker)
            thread.daemon = True
            thread.start()
            self.download_threads.append(thread)
            
    def stop_download(self):
        """停止所有下载任务"""
        self.is_downloading = False
        
        # 清空下载队列
        while not self.download_queue.empty():
            try:
                article = self.download_queue.get_nowait()
                self.download_results[article['url']]['status'] = '已取消'
                if self.progress_callback:
                    self.progress_callback(article['url'], self.download_results[article['url']])
            except:
                pass
                
        # 等待所有线程结束
        for thread in self.download_threads:
            if thread.is_alive():
                thread.join(0.1)
            
        self.download_threads = []
        
    def _download_worker(self):
        """下载工作线程"""
        while self.is_downloading:
            try:
                article = self.download_queue.get(timeout=1)
            except Empty:
                if self.download_queue.empty():
                    time.sleep(0.5)
                    if self.is_downloading:
                        self.is_downloading = False
                break
            except Exception as e:
                print(f"下载队列异常: {str(e)}")
                break
                
            try:
                # 更新状态为下载中
                self.download_results[article['url']]['status'] = '下载中'
                if self.progress_callback:
                    self.progress_callback(article['url'], self.download_results[article['url']])
                
                # 创建下载器实例
                downloader = WeChatArticleDownloader(save_dir=self.save_dir)
                
                # 下载文章
                success, file_path = downloader.download_article(article['url'])
                
                # 更新下载状态
                if success:
                    self.download_results[article['url']] = {
                        'status': '已下载',
                        'file_path': file_path
                    }
                else:
                    self.download_results[article['url']] = {
                        'status': '下载失败',
                        'file_path': None
                    }
                    
                if self.progress_callback:
                    self.progress_callback(article['url'], self.download_results[article['url']])
                    
            except Exception as e:
                self.download_results[article['url']] = {
                    'status': f'下载失败: {str(e)}',
                    'file_path': None
                }
                if self.progress_callback:
                    self.progress_callback(article['url'], self.download_results[article['url']])
                
            finally:
                self.download_queue.task_done()
                
    def get_article_status(self, article_url):
        """获取文章的下载状态"""
        return self.download_results.get(article_url, {'status': '未知', 'file_path': None})
    
    def create_zip_package(self, account_name, include_excel=False, excel_file_path=None):
        """创建压缩包"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            zip_filename = f"{account_name}_{timestamp}.zip"
            
            # 确保zip文件保存在static/exports目录下
            exports_dir = os.path.join('static', 'exports')
            os.makedirs(exports_dir, exist_ok=True)
            zip_path = os.path.join(exports_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # 添加markdown文件夹
                markdown_dir = self.save_dir
                if os.path.exists(markdown_dir):
                    for root, dirs, files in os.walk(markdown_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            # 在压缩包中保持相对路径结构
                            arcname = os.path.join('markdown', os.path.relpath(file_path, markdown_dir))
                            zipf.write(file_path, arcname)
                
                # 如果需要包含Excel文件
                if include_excel and excel_file_path and os.path.exists(excel_file_path):
                    zipf.write(excel_file_path, os.path.basename(excel_file_path))
            
            return zip_path
        except Exception as e:
            print(f"创建压缩包失败: {str(e)}")
            return None
        
if __name__ == "__main__":
    # 测试入口
    # 输入公众号文章链接
    article_url = input("请输入微信公众号文章链接: ").strip()
    # 选择保存格式
    print("请选择保存格式：1-Markdown  2-Word(docx)  3-PDF")
    fmt_choice = input("输入数字选择格式: ").strip()
    fmt_map = {'1': 'md', '2': 'docx', '3': 'pdf'}
    file_format = fmt_map.get(fmt_choice, 'md')

    # 可选：输入公众号名称（用于保存路径分组）
    account_name = input("请输入公众号名称（可选，直接回车跳过）: ").strip() or None

    # 创建下载器实例
    downloader = WeChatArticleDownloader(save_dir="downloads")

    # 下载文章
    success, file_path = downloader.download_article(article_url, file_format=file_format, account_name=account_name)
    if success:
        print(f"文章已成功保存到: {file_path}")
    else:
        print("下载失败，请检查链接或网络。")